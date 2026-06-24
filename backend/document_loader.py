import os
import re
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from pypdf import PdfReader
import docx


# ----------------------------------
# Text Cleaning Function
# ----------------------------------

def clean_text(text: str) -> str:
    """
    Clean extracted text by removing
    extra spaces and unnecessary characters.
    """

    if not text:
        return ""

    # Replace multiple spaces/newlines
    text = re.sub(r"\s+", " ", text)

    # Remove non-printable characters
    text = text.encode("ascii", "ignore").decode()

    return text.strip()


# ----------------------------------
# PDF Loader
# ----------------------------------

def load_pdf(file_path: str):
    """
    Extract text from PDF while
    keeping page information.
    """

    documents = []

    pdf = PdfReader(file_path)

    for page_number, page in enumerate(pdf.pages):

        text = page.extract_text()

        text = clean_text(text)

        if text:

            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": Path(file_path).name,
                        "page": page_number + 1,
                        "type": "pdf"
                    }
                )
            )

    return documents


# ----------------------------------
# DOCX Loader
# ----------------------------------

def load_docx(file_path: str):
    """
    Extract text from DOCX files.
    """

    document = docx.Document(file_path)

    text = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
    )

    text = clean_text(text)

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": Path(file_path).name,
                "page": 1,
                "type": "docx"
            }
        )
    ]


# ----------------------------------
# TXT Loader
# ----------------------------------

def load_txt(file_path: str):
    """
    Load plain text files.
    """

    with open(
        file_path,
        "r",
        encoding="utf-8",
        errors="ignore"
    ) as file:

        text = file.read()

    text = clean_text(text)

    if not text:
        return []

    return [
        Document(
            page_content=text,
            metadata={
                "source": Path(file_path).name,
                "page": 1,
                "type": "txt"
            }
        )
    ]


# ----------------------------------
# Detect File Type
# ----------------------------------

def load_single_document(file_path: str):

    extension = Path(file_path).suffix.lower()

    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".txt": load_txt
    }

    if extension not in loaders:
        raise ValueError(
            f"Unsupported file format: {extension}"
        )

    return loaders[extension](file_path)


# ----------------------------------
# Load All Documents from Folder
# ----------------------------------

def load_documents(folder_path: str):

    all_documents = []

    for filename in os.listdir(folder_path):

        file_path = os.path.join(
            folder_path,
            filename
        )

        if os.path.isfile(file_path):

            try:

                docs = load_single_document(
                    file_path
                )

                all_documents.extend(docs)

                print(
                    f"Loaded: {filename}"
                )

            except Exception as e:

                print(
                    f"Error loading {filename}: {e}"
                )

    return all_documents


# ----------------------------------
# Document Chunking
# ----------------------------------

def split_documents(
    documents,
    chunk_size=700,
    chunk_overlap=100
):
    """
    Split large documents into chunks
    suitable for embedding.
    """

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=[
            "\n\n",
            "\n",
            ". ",
            " ",
            ""
        ]
    )

    chunks = splitter.split_documents(
        documents
    )


    # Add chunk ID metadata
    for idx, chunk in enumerate(chunks):

        chunk.metadata["chunk_id"] = idx + 1


    print(
        f"Created {len(chunks)} chunks"
    )

    return chunks


# ----------------------------------
# Testing the loader directly
# ----------------------------------

if __name__ == "__main__":

    folder = "../data/uploads"

    docs = load_documents(folder)

    chunks = split_documents(
        docs
    )

    print("\nSample Chunk:\n")
    print(
        chunks[0].page_content[:500]
    )

    print("\nMetadata:")
    print(
        chunks[0].metadata
    )